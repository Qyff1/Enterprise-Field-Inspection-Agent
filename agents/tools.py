"""
工具定义模块
包含通用工具 + 外勤行程核验专用工具
每个工具使用 @tool 装饰器注册
"""
import base64
import os
import json
import re
from datetime import datetime
from typing import Dict, Any, Optional

from PIL import Image
from langchain.tools import tool
import requests
import docx
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn

# ============================================================
# 路径常量
# ============================================================
_UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'storage', 'uploads')
_PROJECT_ROOT = os.path.dirname(os.path.dirname(__file__))
_OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'storage', 'outputs')

os.makedirs(_OUTPUT_DIR, exist_ok=True)


# ============================================================
# 通用工具
# ============================================================

@tool
def web_search(query: str, count: int = 10) -> str:
    """
    使用 Bocha Web Search API 进行网页搜索。

    参数:
    - query: 搜索关键词
    - count: 返回的搜索结果数量（默认10）

    返回:
    - 搜索结果的详细信息
    """
    BOCHA_API_KEY = os.getenv("BOCHA_API_KEY", "")
    url = 'https://api.bochaai.com/v1/web-search'
    headers = {
        'Authorization': f'Bearer {BOCHA_API_KEY}',
        'Content-Type': 'application/json'
    }
    data = {"query": query, "freshness": "noLimit", "summary": True, "count": count}

    response = requests.post(url, headers=headers, json=data)
    if response.status_code != 200:
        return f"搜索API请求失败，状态码: {response.status_code}"

    json_response = response.json()
    try:
        if json_response["code"] != 200 or not json_response["data"]:
            return f"搜索失败: {json_response.get('msg', '未知错误')}"

        webpages = json_response["data"]["webPages"]["value"]
        if not webpages:
            return "未找到相关结果。"

        formatted_results = ""
        for idx, page in enumerate(webpages, start=1):
            formatted_results += (
                f"引用: {idx}\n"
                f"标题: {page['name']}\n"
                f"URL: {page['url']}\n"
                f"摘要: {page['summary']}\n"
                f"网站名称: {page['siteName']}\n"
                f"发布时间: {page['dateLastCrawled']}\n\n"
            )
        return formatted_results.strip()
    except Exception as e:
        return f"搜索结果解析失败: {str(e)}"


@tool
def search_knowledge(query: str) -> str:
    """
    使用 RAG 语义搜索本地知识库 (knowledge.md) 和对话记忆 (memory.md)。
    当遇到不确定的问题时，**必须先调用此工具**查找本地知识。

    参数:
    - query: 搜索查询语句

    返回:
    - 知识库中语义相关的条目内容及来源
    """
    try:
        from agents.rag import get_rag
        rag = get_rag()
        result = rag.search(query, top_k=5)
        return result['formatted']
    except ImportError:
        knowledge_path = os.path.join(_PROJECT_ROOT, "data", "knowledge.md")
        memory_path = os.path.join(_PROJECT_ROOT, "data", "memory.md")
        parts = []
        for path, label in [(knowledge_path, "知识库"), (memory_path, "对话记忆")]:
            if os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()[:2000]
                parts.append(f"## {label}\n{content}")
        return '\n\n'.join(parts) if parts else "本地知识库和记忆中暂无内容。"
    except Exception as e:
        return f"知识库搜索失败: {str(e)}"


@tool
def save_memory(title: str, content: str, memory_type: str = "其他") -> str:
    """
    将重要信息保存到项目记忆文件 memory.md 中，用于跨对话记忆持久化。

    参数:
    - title: 记忆标题
    - content: 记忆内容
    - memory_type: 记忆类型（偏好/决策/问题/代码/其他/外勤记录）

    返回:
    - 保存成功或失败的提示信息
    """
    memory_path = os.path.join(_PROJECT_ROOT, "data", "memory.md")
    if not os.path.exists(memory_path):
        return f"记忆文件不存在: {memory_path}"

    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = (
        f"\n### [{datetime.now().strftime('%Y-%m-%d')}] {title}\n"
        f"- **类型**: {memory_type}\n"
        f"- **时间**: {date_str}\n"
        f"- **内容**: {content}\n"
    )

    try:
        with open(memory_path, "r", encoding="utf-8") as f:
            existing = f.read()
        if f"[{datetime.now().strftime('%Y-%m-%d')}] {title}" in existing:
            return f"记忆已存在: {title}"

        with open(memory_path, "a", encoding="utf-8") as f:
            f.write(entry)

        try:
            from agents.rag import get_rag
            rag = get_rag()
            rag.add_entry(
                text=entry,
                title=f"[{datetime.now().strftime('%Y-%m-%d')}] {title}",
                source_type="memory",
                metadata={"type_name": memory_type}
            )
        except Exception:
            pass

        return f"记忆已保存: {title}\n类型: {memory_type}\n内容预览: {content[:100]}..."
    except Exception as e:
        return f"保存记忆失败: {str(e)}"


@tool
def save_knowledge(title: str, content: str, category: str = "通用", tags: str = "") -> str:
    """
    将知识点保存到项目知识库文件 knowledge.md 中。

    参数:
    - title: 知识标题
    - content: 知识内容
    - category: 知识分类（如 审计规则/外勤政策/异常模式）
    - tags: 标签，逗号分隔

    返回:
    - 保存成功或失败的提示信息
    """
    knowledge_path = os.path.join(_PROJECT_ROOT, "data", "knowledge.md")
    if not os.path.exists(knowledge_path):
        return f"知识库文件不存在: {knowledge_path}"

    date_str = datetime.now().strftime("%Y-%m-%d")
    tags_str = tags if tags else category
    entry = (
        f"\n### [{category}] {title}\n"
        f"- **标签**: {tags_str}\n"
        f"- **内容**: {content}\n"
        f"- **来源**: 审计系统自动记录\n"
        f"- **更新**: {date_str}\n"
    )

    try:
        with open(knowledge_path, "r", encoding="utf-8") as f:
            existing = f.read()
        if f"[{category}] {title}" in existing:
            return f"知识条目已存在: {title}"

        with open(knowledge_path, "a", encoding="utf-8") as f:
            f.write(entry)

        try:
            from agents.rag import get_rag
            rag = get_rag()
            rag.add_entry(
                text=entry,
                title=f"[{category}] {title}",
                source_type="knowledge",
                metadata={"category": category, "tags": tags_str}
            )
        except Exception:
            pass

        return f"知识已保存: [{category}] {title}"
    except Exception as e:
        return f"保存知识失败: {str(e)}"


@tool
def read_local_file(file_path: str) -> str:
    """
    读取本地文件内容。支持 docx、txt、md、json、csv、图片等格式。

    参数:
    - file_path: 文件的绝对路径

    返回:
    - 文件内容字符串 | 图片Base64信息 | 失败返回错误字符串
    """
    if not os.path.exists(file_path):
        alt_path = os.path.join(_UPLOAD_DIR, os.path.basename(file_path))
        if os.path.exists(alt_path):
            file_path = alt_path
        else:
            if os.path.isdir(_UPLOAD_DIR):
                for f in os.listdir(_UPLOAD_DIR):
                    if file_path in f or f.endswith(file_path):
                        file_path = os.path.join(_UPLOAD_DIR, f)
                        break
                else:
                    return "读取文件失败：文件不存在"
            else:
                return "读取文件失败：文件不存在"

    suffix = os.path.splitext(file_path)[-1].lower()
    text_suffix = (".txt", ".md", ".json", ".csv", ".py", ".log", ".html")
    img_suffix = (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp")

    try:
        if suffix == ".docx":
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            # 也读取表格内容
            for table in doc.tables:
                content += "\n[表格]\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    content += " | ".join(cells) + "\n"
        elif suffix in text_suffix:
            content = ""
            for encoding in ["utf-8", "gbk", "gb2312", "utf-8-sig", "gb18030", "cp1252", "latin-1"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return "读取文件失败: 无法识别文件编码"
        elif suffix in img_suffix:
            with Image.open(file_path) as img:
                w, h = img.size
            with open(file_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")
            return f"【图片文件】尺寸：{w}*{h} 格式：{suffix}\n图片Base64：{img_b64}"
        else:
            return f"读取文件失败: 暂不支持{suffix}格式文件"

        is_memory_file = os.path.basename(file_path) == "memory.md"
        max_len = 10000 if is_memory_file else 5000
        if len(content) > max_len:
            content = content[:max_len] + "\n...内容过长，已截断"
        return content.strip()
    except Exception as e:
        return f"读取文件失败: {str(e)}"


# ============================================================
# 外勤行程核验专用工具
# ============================================================

# 行程字段关键词匹配表（中英文）
_TRIP_FIELD_PATTERNS = {
    "employee_name": ["员工姓名", "姓名", "员工", "申请人", "name", "employee"],
    "trip_date": ["日期", "出行日期", "出差日期", "date"],
    "departure_time": ["出发时间", "出发时刻", "时间", "time"],
    "origin": ["出发地", "起点", "出发地点", "出发地址", "origin", "from"],
    "destination": ["目的地", "终点", "到达地点", "目的地址", "destination", "to"],
    "reported_mileage": ["里程", "公里数", "申报里程", "距离", "mileage", "distance", "km"],
    "trip_purpose": ["目的", "事由", "出行目的", "出差事由", "purpose", "reason"],
    "vehicle_plate": ["车牌", "车牌号", "车辆", "plate", "vehicle"],
}


def _match_field_name(text: str) -> Optional[str]:
    """将表头文本匹配到标准字段名"""
    text_lower = text.lower().strip()
    for field_name, keywords in _TRIP_FIELD_PATTERNS.items():
        for kw in keywords:
            if kw in text_lower:
                return field_name
    return None


@tool
def extract_trip_data(file_path: str) -> str:
    """
    解析外勤报告文档（docx/xlsx/txt），提取结构化行程数据。

    参数:
    - file_path: 文档文件的绝对路径

    返回:
    - JSON格式的行程数据，包含 meta 和 records
    """
    import json as _json

    if not os.path.exists(file_path):
        alt = os.path.join(_UPLOAD_DIR, os.path.basename(file_path))
        if os.path.exists(alt):
            file_path = alt
        else:
            return _json.dumps({"error": f"文件不存在: {file_path}"}, ensure_ascii=False)

    suffix = os.path.splitext(file_path)[-1].lower()
    records = []
    parse_errors = []

    try:
        if suffix == ".docx":
            doc_obj = docx.Document(file_path)

            # 策略1: 查找表格中的行程数据
            for table in doc_obj.tables:
                rows = table.rows
                if len(rows) < 2:
                    continue

                # 解析表头行，映射列索引到字段名
                header_map = {}
                header_row = rows[0]
                for col_idx, cell in enumerate(header_row.cells):
                    field = _match_field_name(cell.text)
                    if field:
                        header_map[col_idx] = field

                if len(header_map) < 3:  # 至少要有3个有效字段才算行程表
                    continue

                # 解析数据行
                for row in rows[1:]:
                    record = {}
                    for col_idx, field_name in header_map.items():
                        if col_idx < len(row.cells):
                            val = row.cells[col_idx].text.strip()
                            if val:
                                record[field_name] = val

                    # 尝试数值化里程
                    if "reported_mileage" in record:
                        try:
                            mileage_str = re.sub(r'[^\d.]', '', record["reported_mileage"])
                            record["reported_mileage"] = float(mileage_str) if mileage_str else 0
                        except ValueError:
                            record["reported_mileage"] = 0

                    if len(record) >= 3:
                        record["record_id"] = (
                            f"{record.get('employee_name', 'unknown')}_"
                            f"{record.get('trip_date', 'unknown')}_"
                            f"{len(records)}"
                        )
                        records.append(record)

            # 策略2: 如果没有表格，从段落文本中提取
            if not records:
                text = "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()])
                # 用 LLM 辅助解析时返回原始文本
                if text:
                    return _json.dumps({
                        "meta": {
                            "file_name": os.path.basename(file_path),
                            "parse_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "record_count": 0,
                            "parse_method": "fallback_text",
                            "parse_errors": ["文档中未找到结构化表格，请人工检查或使用 read_local_file 查看原始内容"]
                        },
                        "records": [],
                        "raw_text_preview": text[:2000]
                    }, ensure_ascii=False, indent=2)

        elif suffix == ".xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                ws = wb.active

                all_rows = list(ws.iter_rows(values_only=True))
                if len(all_rows) < 2:
                    return _json.dumps({"error": "Excel文件行数不足"}, ensure_ascii=False)

                header_map = {}
                header_row = all_rows[0]
                for col_idx, cell_val in enumerate(header_row):
                    if cell_val:
                        field = _match_field_name(str(cell_val))
                        if field:
                            header_map[col_idx] = field

                if len(header_map) < 3:
                    return _json.dumps({
                        "error": "无法识别Excel表头字段",
                        "headers_found": [str(c) for c in header_row if c]
                    }, ensure_ascii=False)

                for row in all_rows[1:]:
                    record = {}
                    for col_idx, field_name in header_map.items():
                        if col_idx < len(row) and row[col_idx] is not None:
                            val = str(row[col_idx]).strip()
                            if val and val != "None":
                                record[field_name] = val

                    if "reported_mileage" in record:
                        try:
                            record["reported_mileage"] = float(re.sub(r'[^\d.]', '', record["reported_mileage"]))
                        except ValueError:
                            record["reported_mileage"] = 0

                    if len(record) >= 3:
                        record["record_id"] = (
                            f"{record.get('employee_name', 'unknown')}_"
                            f"{record.get('trip_date', 'unknown')}_"
                            f"{len(records)}"
                        )
                        records.append(record)

                wb.close()
            except ImportError:
                return _json.dumps({"error": "openpyxl 未安装，无法解析 xlsx 文件"}, ensure_ascii=False)

        elif suffix in (".txt", ".md", ".csv"):
            content = ""
            for enc in ["utf-8", "gbk", "utf-8-sig", "gb2312"]:
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if not content:
                return _json.dumps({"error": "无法读取文件编码"}, ensure_ascii=False)

            # CSV 尝试
            if suffix == ".csv":
                lines = content.strip().split("\n")
                if len(lines) >= 2:
                    headers = [h.strip() for h in lines[0].split(",")]
                    header_map = {}
                    for col_idx, h in enumerate(headers):
                        field = _match_field_name(h)
                        if field:
                            header_map[col_idx] = field

                    if len(header_map) >= 3:
                        for line in lines[1:]:
                            cells = [c.strip() for c in line.split(",")]
                            record = {}
                            for col_idx, field_name in header_map.items():
                                if col_idx < len(cells) and cells[col_idx]:
                                    record[field_name] = cells[col_idx]

                            if "reported_mileage" in record:
                                try:
                                    record["reported_mileage"] = float(re.sub(r'[^\d.]', '', record["reported_mileage"]))
                                except ValueError:
                                    record["reported_mileage"] = 0

                            if len(record) >= 3:
                                record["record_id"] = (
                                    f"{record.get('employee_name', 'unknown')}_"
                                    f"{record.get('trip_date', 'unknown')}_"
                                    f"{len(records)}"
                                )
                                records.append(record)

            if not records:
                return _json.dumps({
                    "meta": {
                        "file_name": os.path.basename(file_path),
                        "parse_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "record_count": 0,
                        "parse_method": "fallback_text",
                        "parse_errors": ["纯文本格式无法自动提取，请使用 read_local_file 查看后手动提取"]
                    },
                    "records": [],
                    "raw_text_preview": content[:2000]
                }, ensure_ascii=False, indent=2)

        else:
            return _json.dumps({"error": f"不支持的文件格式: {suffix}"}, ensure_ascii=False)

    except Exception as e:
        parse_errors.append(str(e))

    # 检查必填字段
    for rec in records:
        missing = []
        for field in ["employee_name", "trip_date", "origin", "destination"]:
            if field not in rec or not rec[field]:
                missing.append(field)
        if missing:
            parse_errors.append(f"记录 {rec.get('record_id', '?')} 缺少字段: {', '.join(missing)}")

    result = {
        "meta": {
            "file_name": os.path.basename(file_path),
            "parse_time": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "record_count": len(records),
            "parse_errors": parse_errors
        },
        "records": records
    }

    return _json.dumps(result, ensure_ascii=False, indent=2)


@tool
def geocode_address(address: str) -> str:
    """
    使用高德地图MCP将地址转换为坐标和标准化名称。

    参数:
    - address: 原始地址字符串

    返回:
    - JSON: {original, standardized, lng, lat, adcode, status}
    """
    import json as _json
    from agents.mcp_services import call_mcp_tool

    # 尝试通过 MCP 调用高德地理编码
    result = call_mcp_tool("maps-geocode", address=address)

    if result:
        try:
            if isinstance(result, str):
                result = _json.loads(result)
            return _json.dumps({
                "original": address,
                "standardized": result.get("formatted_address", result.get("address", address)),
                "lng": result.get("lng", result.get("location", {}).get("lng", 0)),
                "lat": result.get("lat", result.get("location", {}).get("lat", 0)),
                "adcode": result.get("adcode", ""),
                "status": "success"
            }, ensure_ascii=False, indent=2)
        except Exception as e:
            pass  # 解析失败回退到 REST API

    # MCP 调用失败，降级到高德 REST API
    from agents.mcp_services import get_amap_api_key
    api_key = get_amap_api_key()
    if not api_key:
        return _json.dumps({
            "original": address, "standardized": address,
            "lng": 0, "lat": 0, "status": "geocode_failed",
            "error": "高德 API Key 不可用"
        }, ensure_ascii=False, indent=2)

    try:
        url = "https://restapi.amap.com/v3/geocode/geo"
        resp = requests.get(url, params={
            "key": api_key, "address": address, "output": "JSON"
        }, timeout=10)

        data = resp.json()
        if data.get("status") == "1" and data.get("geocodes"):
            geo = data["geocodes"][0]
            location = geo.get("location", "0,0").split(",")
            return _json.dumps({
                "original": address,
                "standardized": geo.get("formatted_address", address),
                "lng": float(location[0]) if len(location) >= 2 else 0,
                "lat": float(location[1]) if len(location) >= 2 else 0,
                "adcode": geo.get("adcode", ""),
                "status": "success_via_rest"
            }, ensure_ascii=False, indent=2)
        else:
            return _json.dumps({
                "original": address, "standardized": address,
                "lng": 0, "lat": 0, "status": "geocode_failed",
                "error": "地址无法识别，建议补充详细地址信息"
            }, ensure_ascii=False, indent=2)
    except Exception as e:
        return _json.dumps({
            "original": address, "standardized": address,
            "lng": 0, "lat": 0, "status": "geocode_error",
            "error": str(e)
        }, ensure_ascii=False, indent=2)


@tool
def calculate_route_mileage(origin: str, destination: str) -> str:
    """
    使用高德地图计算两地之间的真实驾车里程和预计时间。

    参数:
    - origin: 出发地地址
    - destination: 目的地地址

    返回:
    - JSON: {origin, destination, distance_km, duration_min, route_summary, status}
    """
    import json as _json
    from agents.mcp_services import call_mcp_tool, get_amap_api_key

    # 尝试通过 MCP 调用高德驾车路线规划
    result = call_mcp_tool("maps-direction-driving", origin=origin, destination=destination)

    if result:
        try:
            if isinstance(result, str):
                result = _json.loads(result)
            route = result.get("route", {})
            paths = route.get("paths", [])
            if paths:
                path = paths[0]
                return _json.dumps({
                    "origin": origin,
                    "destination": destination,
                    "distance_km": round(path.get("distance", 0) / 1000, 1),
                    "duration_min": round(path.get("duration", 0) / 60, 0),
                    "route_summary": "高德推荐路线",
                    "status": "success"
                }, ensure_ascii=False, indent=2)
        except Exception:
            pass  # 解析失败回退到 REST API

    # MCP 失败，降级到 REST API
    api_key = get_amap_api_key()
    if not api_key:
        return _json.dumps({
            "origin": origin, "destination": destination,
            "distance_km": 0, "duration_min": 0,
            "status": "amap_unavailable", "error": "高德 API Key 不可用"
        }, ensure_ascii=False, indent=2)

    try:
        # Step 1: 地理编码出发地和目的地
        geo_url = "https://restapi.amap.com/v3/geocode/geo"
        origin_resp = requests.get(geo_url, params={
            "key": api_key, "address": origin, "output": "JSON"
        }, timeout=10).json()
        dest_resp = requests.get(geo_url, params={
            "key": api_key, "address": destination, "output": "JSON"
        }, timeout=10).json()

        if origin_resp.get("status") != "1" or not origin_resp.get("geocodes"):
            return _json.dumps({
                "origin": origin, "destination": destination,
                "distance_km": 0, "duration_min": 0,
                "status": "geocode_failed",
                "error": f"出发地无法识别: {origin}"
            }, ensure_ascii=False, indent=2)

        if dest_resp.get("status") != "1" or not dest_resp.get("geocodes"):
            return _json.dumps({
                "origin": origin, "destination": destination,
                "distance_km": 0, "duration_min": 0,
                "status": "geocode_failed",
                "error": f"目的地无法识别: {destination}"
            }, ensure_ascii=False, indent=2)

        origin_loc = origin_resp["geocodes"][0]["location"]
        dest_loc = dest_resp["geocodes"][0]["location"]

        # Step 2: 驾车路线规划
        dir_url = "https://restapi.amap.com/v3/direction/driving"
        dir_resp = requests.get(dir_url, params={
            "key": api_key,
            "origin": origin_loc,
            "destination": dest_loc,
            "extensions": "base",
            "output": "JSON"
        }, timeout=10).json()

        if dir_resp.get("status") == "1" and dir_resp.get("route", {}).get("paths"):
            path = dir_resp["route"]["paths"][0]
            return _json.dumps({
                "origin": origin,
                "destination": destination,
                "origin_coord": origin_loc,
                "dest_coord": dest_loc,
                "distance_km": round(int(path.get("distance", 0)) / 1000, 1),
                "duration_min": round(int(path.get("duration", 0)) / 60, 0),
                "route_summary": f"高德驾车路线 ({origin_resp['geocodes'][0].get('formatted_address', origin)} → {dest_resp['geocodes'][0].get('formatted_address', destination)})",
                "status": "success_via_rest"
            }, ensure_ascii=False, indent=2)
        else:
            return _json.dumps({
                "origin": origin, "destination": destination,
                "distance_km": 0, "duration_min": 0,
                "status": "route_failed",
                "error": "无法规划驾车路线"
            }, ensure_ascii=False, indent=2)

    except Exception as e:
        return _json.dumps({
            "origin": origin, "destination": destination,
            "distance_km": 0, "duration_min": 0,
            "status": "error", "error": str(e)
        }, ensure_ascii=False, indent=2)


@tool
def compare_mileage(reported: float, actual: float) -> str:
    """
    比对申报里程与实际里程，计算偏差并给出状态标记。

    参数:
    - reported: 申报里程（公里）
    - actual: 高德地图实际里程（公里）

    返回:
    - JSON: {reported, actual, discrepancy_km, discrepancy_pct, status, over_reported}
    """
    import json as _json
    from config import VERIFY_CONFIG

    if actual <= 0:
        return _json.dumps({
            "reported_km": reported,
            "actual_km": actual,
            "discrepancy_km": reported,
            "discrepancy_pct": 100.0,
            "status": "unverified",
            "over_reported": True,
            "note": "实际里程为0或无效，无法比对"
        }, ensure_ascii=False, indent=2)

    threshold = VERIFY_CONFIG.get("mileage_threshold_pct", 15)
    warn_pct = VERIFY_CONFIG.get("mileage_warn_pct", 5)

    discrepancy = reported - actual
    pct = abs(discrepancy) / actual * 100

    if pct > threshold:
        status = "red"
    elif pct > warn_pct:
        status = "yellow"
    else:
        status = "green"

    return _json.dumps({
        "reported_km": reported,
        "actual_km": actual,
        "discrepancy_km": round(discrepancy, 1),
        "discrepancy_pct": round(pct, 2),
        "status": status,
        "over_reported": discrepancy > 0
    }, ensure_ascii=False, indent=2)


@tool
def check_historical_patterns(employee: str, location: str) -> str:
    """
    查询员工历史出行模式，检测是否为常规行程。

    参数:
    - employee: 员工姓名
    - location: 目的地地址

    返回:
    - JSON: {employee, location, visit_count, is_regular, historical_mileage_avg, patterns}
    """
    import json as _json

    try:
        from agents.rag import get_rag
        rag = get_rag()

        # 构造查询：员工 + 地点 + 外勤
        query = f"{employee} {location} 外勤 行程"
        result = rag.search(query, top_k=10)

        if not result.get("found"):
            return _json.dumps({
                "employee": employee,
                "location": location,
                "visit_count": 0,
                "is_regular": False,
                "historical_mileage_avg": 0,
                "first_seen": None,
                "patterns": [],
                "note": "未找到该员工的历史出行记录"
            }, ensure_ascii=False, indent=2)

        # 解析检索结果，提取里程和时间信息
        raw_entries = result.get("raw", [])
        visits = []
        mileages = []
        for entry in raw_entries:
            title = entry.get("title", "")
            if employee in title or location in title:
                visits.append({"title": title, "relevance": entry.get("relevance", 0)})

        # 从 formatted 内容中提取里程数值
        formatted = result.get("formatted", "")
        mileage_matches = re.findall(r'里程[：:]\s*(\d+\.?\d*)\s*(?:km|公里)?', formatted)
        for m in mileage_matches:
            try:
                mileages.append(float(m))
            except ValueError:
                pass

        visit_count = len(visits) + len(mileage_matches)
        from config import VERIFY_CONFIG
        threshold = VERIFY_CONFIG.get("regular_location_threshold", 3)
        is_regular = visit_count >= threshold

        avg_mileage = round(sum(mileages) / len(mileages), 1) if mileages else 0

        return _json.dumps({
            "employee": employee,
            "location": location,
            "visit_count": visit_count,
            "is_regular": is_regular,
            "historical_mileage_avg": avg_mileage,
            "first_seen": visits[0]["title"] if visits else None,
            "patterns": visits[:5],
            "note": "常规地点" if is_regular else "首次或低频访问地点，需关注"
        }, ensure_ascii=False, indent=2)

    except ImportError:
        return _json.dumps({
            "employee": employee, "location": location,
            "visit_count": 0, "is_regular": False,
            "historical_mileage_avg": 0, "patterns": [],
            "note": "RAG 知识库不可用，无法查询历史模式"
        }, ensure_ascii=False, indent=2)
    except Exception as e:
        return _json.dumps({
            "employee": employee, "location": location,
            "visit_count": 0, "is_regular": False,
            "error": str(e)
        }, ensure_ascii=False, indent=2)


@tool
def generate_audit_report(findings_json: str) -> str:
    """
    汇总所有核验结果，生成结构化Markdown审计报告。

    参数:
    - findings_json: JSON字符串，包含所有核验步骤的输出

    返回:
    - Markdown格式的完整审计报告
    """
    import json as _json

    try:
        findings = _json.loads(findings_json) if isinstance(findings_json, str) else findings_json
    except _json.JSONDecodeError:
        return f"# 审计报告生成失败\n\n无法解析输入数据，请检查JSON格式。\n\n原始输入:\n```\n{findings_json[:500]}\n```"

    meta = findings.get("meta", {})
    records = findings.get("records", [])
    anomalies = findings.get("anomalies", [])

    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    report = f"""# 🔍 外勤行程核验审计报告

## 一、基本信息

| 项目 | 内容 |
|------|------|
| 审计日期 | {now} |
| 审计文件 | {meta.get('file_name', '未知')} |
| 行程总数 | {meta.get('record_count', len(records))} 条 |
| 异常数量 | {len([a for a in anomalies if isinstance(a, dict) and a.get('overall_risk_score', 0) >= 0.5])} 条 |
| 核验工具 | 企业外勤核验智能Agent v1.0（高德地图MCP） |

## 二、核验结论

"""

    # 统计状态
    green_count = 0
    yellow_count = 0
    red_count = 0
    for rec in records:
        mileage_result = rec.get("mileage_comparison", {})
        status = mileage_result.get("status", "unknown")
        if status == "green":
            green_count += 1
        elif status == "yellow":
            yellow_count += 1
        elif status == "red":
            red_count += 1

    total_risk = 0
    anomaly_count = 0
    for a in anomalies:
        if isinstance(a, dict):
            score = a.get("overall_risk_score", 0)
            if score >= 0.5:
                anomaly_count += 1
            total_risk = max(total_risk, score)

    if total_risk >= 0.8:
        risk_level = "🔴 高风险"
    elif total_risk >= 0.5:
        risk_level = "🟡 中风险"
    else:
        risk_level = "🟢 低风险"

    report += f"""| 统计项 | 数量 |
|--------|------|
| 🟢 正常通过 | {green_count} 条 |
| 🟡 预警关注 | {yellow_count} 条 |
| 🔴 异常需复核 | {red_count} 条 |
| **整体风险等级** | **{risk_level}** |

## 三、逐条核验详情

"""

    for i, rec in enumerate(records):
        report += f"### 记录{i + 1}: {rec.get('employee_name', '?')} {rec.get('trip_date', '?')} {rec.get('origin', '?')} → {rec.get('destination', '?')}\n\n"

        geo_origin = rec.get("geo_origin", {})
        geo_dest = rec.get("geo_dest", {})
        mileage = rec.get("mileage_comparison", {})
        hist = rec.get("historical_check", {})

        status_emoji = {"green": "🟢", "yellow": "🟡", "red": "🔴", "unverified": "⚪"}.get(
            mileage.get("status", "unknown"), "⚪"
        )

        report += f"| 项目 | 申报值 | 核验值 | 状态 |\n"
        report += f"|------|--------|--------|------|\n"
        report += f"| 出发地 | {rec.get('origin', '?')} | {geo_origin.get('standardized', '未核验')} | ✅ |\n"
        report += f"| 目的地 | {rec.get('destination', '?')} | {geo_dest.get('standardized', '未核验')} | ✅ |\n"
        reported_km = rec.get('reported_mileage', 0)
        actual_km = mileage.get('actual_km', 0)
        report += f"| 里程 | {reported_km}km | {actual_km}km | {status_emoji} 偏差{mileage.get('discrepancy_pct', 0)}% |\n"
        report += f"| 预计耗时 | - | {mileage.get('duration_min', '?')}分钟 | - |\n"

        # 历史检测
        is_regular = hist.get("is_regular", False)
        visit_count = hist.get("visit_count", 0)
        report += f"\n- **历史模式**: {'✅ 常规地点' if is_regular else '⚠️ 新地点/低频'} (历史访问 {visit_count} 次)"
        report += f"\n- **出行目的**: {rec.get('trip_purpose', '未填写')}"
        report += f"\n- **车牌号**: {rec.get('vehicle_plate', '未填写')}"

        # 该记录的异常
        rec_anomalies = [a for a in anomalies if isinstance(a, dict) and a.get("record_id") == rec.get("record_id")]
        if rec_anomalies:
            for anom_list in rec_anomalies:
                items = anom_list.get("anomalies", [anom_list])
                for a in items:
                    if isinstance(a, dict) and a.get("confidence", 0) >= 0.5:
                        conf = a.get("confidence", 0)
                        flag = "🔴" if conf >= 0.8 else "🟡"
                        report += f"\n- {flag} **异常**: {a.get('description', '')} (置信度: {conf:.0%})"

        report += "\n\n---\n\n"

    # 异常汇总
    high_risk = [a for a in anomalies if isinstance(a, dict) and a.get("overall_risk_score", 0) >= 0.8]
    mid_risk = [a for a in anomalies if isinstance(a, dict) and 0.5 <= a.get("overall_risk_score", 0) < 0.8]

    if high_risk or mid_risk:
        report += "## 四、异常汇总\n\n"

        if high_risk:
            report += "### 🔴 高风险项（需人工复核）\n\n"
            for a in high_risk:
                report += f"- **记录**: {a.get('record_id', '?')}\n"
                for item in a.get("anomalies", []):
                    report += f"  - {item.get('type', '?')}: {item.get('description', '?')} (置信度: {item.get('confidence', 0):.0%})\n"
                    report += f"    证据: {item.get('evidence', '无')}\n"
                report += f"  - **处理建议**: {a.get('recommendation', '建议人工核查')}\n\n"

        if mid_risk:
            report += "### 🟡 中风险项（建议关注）\n\n"
            for a in mid_risk:
                report += f"- **记录**: {a.get('record_id', '?')}\n"
                for item in a.get("anomalies", []):
                    report += f"  - {item.get('type', '?')}: {item.get('description', '?')} (置信度: {item.get('confidence', 0):.0%})\n"
                report += "\n"

    report += "## 五、系统说明\n\n"
    report += f"""| 项目 | 说明 |
|------|------|
| 地理数据来源 | 高德地图 MCP 服务 |
| 历史数据来源 | memory.md（RAG 向量检索） |
| 里程偏差红色阈值 | 15% |
| 里程偏差黄色阈值 | 5% |
| 时间冲突判定 | 30分钟内 |
| 常规地点判定 | 历史访问 ≥ 3 次 |
| 报告生成时间 | {now} |
| 核验工具版本 | 企业外勤核验智能Agent v1.0 |
"""

    return report


@tool
def annotate_document(file_path: str, issues_json: str) -> str:
    """
    对原始 docx 文档添加高亮标注，标记异常行程数据。

    参数:
    - file_path: 原始 docx 文件路径
    - issues_json: JSON字符串，异常项列表 [{record_id, field, status, note}]

    返回:
    - 标注后的文件保存路径
    """
    import json as _json
    import shutil

    try:
        issues = _json.loads(issues_json) if isinstance(issues_json, str) else issues_json
    except _json.JSONDecodeError:
        return f"标注失败: 无法解析异常数据 JSON"

    if not os.path.exists(file_path):
        alt = os.path.join(_UPLOAD_DIR, os.path.basename(file_path))
        if os.path.exists(alt):
            file_path = alt
        else:
            return f"标注失败: 文件不存在 {file_path}"

    if not file_path.lower().endswith(".docx"):
        return "标注失败: 仅支持 docx 格式文件的标注"

    try:
        # 复制文件到输出目录
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_name = f"{base_name}_audited_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(_OUTPUT_DIR, output_name)
        shutil.copy2(file_path, output_path)

        # 打开副本进行标注
        doc_obj = docx.Document(output_path)

        for issue in issues:
            status = issue.get("status", "yellow")
            note = issue.get("note", "")
            highlight_color = "RED" if status == "red" else "YELLOW"

            # 在文档末尾添加审计批注页
            doc_obj.add_paragraph("─" * 40)
            p = doc_obj.add_paragraph()
            run = p.add_run(f"[审计标注] ")
            run.bold = True
            run.font.size = Pt(10)

            label_run = p.add_run(
                f"记录: {issue.get('record_id', '?')} | "
                f"字段: {issue.get('field', '?')} | "
                f"状态: {status.upper()} | "
                f"说明: {note}"
            )
            label_run.font.size = Pt(10)
            if status == "red":
                label_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                label_run.font.color.rgb = RGBColor(255, 165, 0)

        doc_obj.save(output_path)
        return _json.dumps({
            "success": True,
            "output_path": output_path,
            "output_name": output_name,
            "issues_annotated": len(issues),
            "note": "已在新文件中添加审计批注（文末），原始文件未修改"
        }, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"标注文档失败: {str(e)}"


# ============================================================
# 工具注册表
# ============================================================
TOOLS = [
    # 通用工具
    search_knowledge,
    web_search,
    read_local_file,
    save_memory,
    save_knowledge,
    # 外勤核验专用工具
    extract_trip_data,
    geocode_address,
    calculate_route_mileage,
    compare_mileage,
    check_historical_patterns,
    generate_audit_report,
    annotate_document,
]


def get_tools_metadata() -> list:
    """获取所有已注册工具的元数据信息（供前端展示）"""
    metadata = []
    for t in TOOLS:
        name = getattr(t, 'name', str(t))
        description = getattr(t, 'description', "无描述")
        short_desc = description.split('\n')[0] if description else "无描述"

        args_schema_str = "无参数"
        args_schema = getattr(t, 'args_schema', None)
        if args_schema:
            try:
                args_schema_str = str(args_schema.schema())
            except Exception:
                args_schema_str = "无法解析参数"

        metadata.append({
            "name": name,
            "description": short_desc,
            "full_description": description,
            "args_schema": args_schema_str,
        })
    return metadata
